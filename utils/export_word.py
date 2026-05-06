from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def export_analyse_word(titre_analyse, contenu_texte, nom_client="", exercice=""):
    doc = Document()
    
    # --- STYLE DU DOCUMENT ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)

    # --- EN-TÊTE (Placeholder pour le logo) ---
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    # NOTE : Quand vous aurez le logo, on ajoutera ici la ligne pour l'image
    run_header = p.add_run("SMD CONSULTING | Superviseur IA")
    run_header.font.color.rgb = RGBColor(31, 119, 180) # Bleu corporate
    run_header.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- TITRE DU RAPPORT ---
    doc.add_paragraph("\n") # Espacement haut
    t = doc.add_heading(titre_analyse, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- CARTOUCHE INFOS CLIENT ---
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = f"Client : {nom_client if nom_client else 'Non spécifié'}"
    cells[1].text = f"Exercice : {exercice if exercice else 'N/A'}"
    
    doc.add_paragraph("\n") # Espacement

    # --- CORPS DU RAPPORT ---
    # On nettoie et on structure le texte provenant de l'IA
    for line in contenu_texte.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('###'):
            # Titre de niveau 3
            h = doc.add_heading(line.replace('###', '').strip(), level=2)
        elif line.startswith('##'):
            # Titre de niveau 2
            h = doc.add_heading(line.replace('##', '').strip(), level=1)
        elif line.startswith('**') and line.endswith('**'):
            # Ligne entièrement en gras (souvent des sous-titres)
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '').strip()).bold = True
        else:
            # Texte normal (on nettoie les éventuels ** au milieu du texte)
            p = doc.add_paragraph(line.replace('**', ''))

    # --- PIED DE PAGE ---
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.text = "Rapport généré par le Superviseur IA SMD Consulting - Confidentiel"
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sauvegarde en buffer pour Streamlit
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
